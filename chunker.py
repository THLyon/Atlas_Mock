# import os
# import json
# import re
# import nltk
# import tiktoken
# import unicodedata
# from docx import Document
# from sentence_transformers import SentenceTransformer, util
# from typing import List

# nltk.download("punkt")

# # Token counter (OpenAI-style)
# enc = tiktoken.encoding_for_model("gpt-3.5-turbo")

# # Semantic model (can upgrade to fine-tuned later)
# model = SentenceTransformer("all-MiniLM-L6-v2")


# def num_tokens(text: str) -> int:
#     return len(enc.encode(text))


# def clean_text(text: str) -> str:
#     if not text:
#         return ""
#     text = text.replace("\u00a0", " ")  # non-breaking space
#     text = text.replace("\ufeff", "")  # BOM
#     text = text.replace("\t", " ")
#     text = text.replace("\u201c", '"').replace("\u201d", '"')
#     text = text.replace("\u2018", "'").replace("\u2019", "'")
#     text = text.replace("\u2013", "-").replace("\u2014", "-")
#     text = unicodedata.normalize("NFKC", text)
#     text = re.sub(r"\s+", " ", text)
#     return text.strip()


# def load_docx(file_path: str) -> str:
#     doc = Document(file_path)
#     return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])


# def semantic_chunk_sentences(text: str, sim_threshold=0.7, max_chunk_tokens=300):
#     paragraphs = [p for p in text.split("\n") if p.strip()]
#     sentence_chunks = []
#     sentence_id = 0
#     para_id = 0
#     para_to_sentences = {}

#     for para in paragraphs:
#         raw_sentences = nltk.sent_tokenize(clean_text(para))
#         if not raw_sentences:
#             continue

#         current_group = []
#         group_sent_ids = []

#         for i, sent in enumerate(raw_sentences):
#             s_clean = clean_text(sent)
#             s_emb = model.encode(s_clean, convert_to_tensor=True)

#             if not current_group:
#                 current_group = [s_clean]
#                 group_sent_ids = [sentence_id]
#                 sentence_chunks.append({
#                     "id": sentence_id,
#                     "text": s_clean,
#                     "tokens": num_tokens(s_clean),
#                     "parentParagraphId": para_id
#                 })
#                 sentence_id += 1
#                 continue

#             prev_emb = model.encode(current_group[-1], convert_to_tensor=True)
#             sim_score = util.pytorch_cos_sim(prev_emb, s_emb).item()

#             if sim_score < sim_threshold or num_tokens(" ".join(current_group + [s_clean])) > max_chunk_tokens:
#                 para_to_sentences[para_id] = group_sent_ids
#                 para_id += 1
#                 current_group = [s_clean]
#                 group_sent_ids = [sentence_id]
#             else:
#                 current_group.append(s_clean)
#                 group_sent_ids.append(sentence_id)

#             sentence_chunks.append({
#                 "id": sentence_id,
#                 "text": s_clean,
#                 "tokens": num_tokens(s_clean),
#                 "parentParagraphId": para_id
#             })
#             sentence_id += 1

#         if group_sent_ids:
#             para_to_sentences[para_id] = group_sent_ids
#             para_id += 1

#     return sentence_chunks, para_to_sentences


# def build_paragraphs(sentences, para_to_sentences, max_tokens_per_section=750):
#     paragraphs = []
#     paragraph_id = 0
#     section_id = 0
#     section_map = {}
#     section_paragraphs = []
#     current_tokens = 0

#     for para_idx, sent_ids in para_to_sentences.items():
#         para_sentences = [s for s in sentences if s["id"] in sent_ids]
#         para_text = " ".join([s["text"] for s in para_sentences])
#         para_tokens = num_tokens(para_text)

#         for sid in sent_ids:
#             for s in sentences:
#                 if s["id"] == sid:
#                     s["parentSectionId"] = section_id

#         paragraphs.append({
#             "id": paragraph_id,
#             "text": para_text,
#             "tokens": para_tokens,
#             "parentSectionId": section_id
#         })

#         section_paragraphs.append(paragraph_id)
#         current_tokens += para_tokens
#         paragraph_id += 1

#         if current_tokens > max_tokens_per_section:
#             section_map[section_id] = section_paragraphs
#             section_id += 1
#             section_paragraphs = []
#             current_tokens = 0

#     if section_paragraphs:
#         section_map[section_id] = section_paragraphs

#     return paragraphs, section_map


# def build_sections(paragraphs, section_map):
#     sections = []
#     for section_id, para_ids in section_map.items():
#         section_text = " ".join([paragraphs[i]["text"] for i in para_ids])
#         token_count = num_tokens(section_text)
#         sections.append({
#             "id": section_id,
#             "text": section_text,
#             "tokens": token_count
#         })
#     return sections


# def save_chunked_output(sentences, paragraphs, sections, output_file="chunk_output.json"):
#     with open(output_file, "w") as f:
#         json.dump({
#             "sentences": sentences,
#             "paragraphs": paragraphs,
#             "sections": sections
#         }, f, indent=2, ensure_ascii=False)
#     print(f"Saved output to {output_file}")


# if __name__ == "__main__":
#     import sys
#     if len(sys.argv) != 2:
#         print("Usage: python chunker.py <path_to_docx>")
#         exit(1)

#     file_path = sys.argv[1]
#     raw_text = load_docx(file_path)
#     print(f"📄 Loaded {len(raw_text)} characters from {file_path}")

#     try:
#         sentences, para_to_sentences = semantic_chunk_sentences(raw_text)
#         paragraphs, section_map = build_paragraphs(sentences, para_to_sentences)
#         sections = build_sections(paragraphs, section_map)
#         save_chunked_output(sentences, paragraphs, sections)
#     except Exception as e:
#         print("Error during chunking:", e)


import os
import json
import re
import nltk
import tiktoken
import unicodedata
from docx import Document
from sentence_transformers import SentenceTransformer, util
import sys
sys.stderr = sys.stdout  # to surface Python errors to Node

nltk.download("punkt")

enc = tiktoken.encoding_for_model("gpt-3.5-turbo")
model = SentenceTransformer("all-MiniLM-L6-v2")

def num_tokens(text: str) -> int:
    return len(enc.encode(text))

def clean_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\u00a0", " ")
    text = text.replace("\ufeff", "")
    text = text.replace("\t", " ")
    text = text.replace("\u201c", '"').replace("\u201d", '"')
    text = text.replace("\u2018", "'").replace("\u2019", "'")
    text = text.replace("\u2013", "-").replace("\u2014", "-")
    text = unicodedata.normalize("NFKC", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()

def load_docx(file_path: str) -> str:
    doc = Document(file_path)
    return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

def semantic_chunk_sentences(text: str, sim_threshold=0.7, max_chunk_tokens=300):
    paragraphs = [p for p in text.split("\n") if p.strip()]
    sentence_chunks = []
    sentence_id = 0
    para_id = 0
    para_to_sentences = {}

    for para in paragraphs:
        raw_sentences = nltk.sent_tokenize(clean_text(para))
        if not raw_sentences:
            continue

        current_group = []
        group_sent_ids = []

        for i, sent in enumerate(raw_sentences):
            s_clean = clean_text(sent)
            s_emb = model.encode(s_clean, convert_to_tensor=True)

            if not current_group:
                current_group = [s_clean]
                group_sent_ids = [sentence_id]
                sentence_chunks.append({
                    "id": f"sentence-{sentence_id}",
                    "type": "sentence",
                    "text": s_clean,
                    "tokens": num_tokens(s_clean),
                    "parentParagraphId": f"paragraph-{para_id}"
                })
                sentence_id += 1
                continue

            prev_emb = model.encode(current_group[-1], convert_to_tensor=True)
            sim_score = util.pytorch_cos_sim(prev_emb, s_emb).item()

            if sim_score < sim_threshold or num_tokens(" ".join(current_group + [s_clean])) > max_chunk_tokens:
                para_to_sentences[f"paragraph-{para_id}"] = group_sent_ids
                para_id += 1
                current_group = [s_clean]
                group_sent_ids = [sentence_id]
            else:
                current_group.append(s_clean)
                group_sent_ids.append(sentence_id)

            sentence_chunks.append({
                "id": f"sentence-{sentence_id}",
                "type": "sentence",
                "text": s_clean,
                "tokens": num_tokens(s_clean),
                "parentParagraphId": f"paragraph-{para_id}"
            })
            sentence_id += 1

        if group_sent_ids:
            para_to_sentences[f"paragraph-{para_id}"] = group_sent_ids
            para_id += 1

    return sentence_chunks, para_to_sentences

def build_paragraphs(sentences, para_to_sentences, max_tokens_per_section=750):
    paragraphs = []
    paragraph_id = 0
    section_id = 0
    section_map = {}
    section_paragraphs = []
    current_tokens = 0

    for para_key, sent_ids in para_to_sentences.items():
        para_sentences = [s for s in sentences if s["id"].startswith("sentence-") and int(s["id"].split("-")[1]) in sent_ids]
        para_text = " ".join([s["text"] for s in para_sentences])
        para_tokens = num_tokens(para_text)

        for s in para_sentences:
            s["parentSectionId"] = f"section-{section_id}"

        paragraphs.append({
            "id": f"paragraph-{paragraph_id}",
            "type": "paragraph",
            "text": para_text,
            "tokens": para_tokens,
            "parentSectionId": f"section-{section_id}"
        })

        section_paragraphs.append(paragraph_id)
        current_tokens += para_tokens
        paragraph_id += 1

        if current_tokens > max_tokens_per_section:
            section_map[f"section-{section_id}"] = section_paragraphs
            section_id += 1
            section_paragraphs = []
            current_tokens = 0

    if section_paragraphs:
        section_map[f"section-{section_id}"] = section_paragraphs

    return paragraphs, section_map

def build_sections(paragraphs, section_map):
    sections = []
    for section_id, para_ids in section_map.items():
        section_text = " ".join([paragraphs[int(i)]["text"] for i in para_ids])
        token_count = num_tokens(section_text)
        sections.append({
            "id": section_id,
            "type": "section",
            "text": section_text,
            "tokens": token_count
        })
    return sections

def save_flat_chunks(*levels, output_file="flattened_chunks.json"):
    all_chunks = []
    for level in levels:
        all_chunks.extend(level)
    with open(output_file, "w") as f:
        json.dump(all_chunks, f, indent=2, ensure_ascii=False)
    print(f"Saved flattened chunks to {output_file}")

if __name__ == "__main__":
    import sys
    if len(sys.argv) != 2:
        print("Usage: python chunker.py <path_to_docx>")
        exit(1)

    file_path = sys.argv[1]
    raw_text = load_docx(file_path)
    print(f"Loaded {len(raw_text)} characters from {file_path}")

    try:
        sentences, para_to_sentences = semantic_chunk_sentences(raw_text)
        paragraphs, section_map = build_paragraphs(sentences, para_to_sentences)
        sections = build_sections(paragraphs, section_map)
        save_flat_chunks(sentences, paragraphs, sections)
    except Exception as e:
        print("Error during chunking:", e)
